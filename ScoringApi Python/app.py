import os
import zipfile
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from flasgger import Swagger, swag_from
from werkzeug.utils import secure_filename
from docx import Document
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import requests
from langdetect import detect, DetectorFactory
import spacy
from sentence_transformers import SentenceTransformer, util
import json

app = Flask(__name__)
CORS(app)
swagger = Swagger(app)

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ZIP_FOLDER'] = 'zips'
app.config['EXTRACTED_FOLDER'] = 'extracted'
app.config['PROCESSED_RECORD'] = 'processed.json'

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])
if not os.path.exists(app.config['ZIP_FOLDER']):
    os.makedirs(app.config['ZIP_FOLDER'])
if not os.path.exists(app.config['EXTRACTED_FOLDER']):
    os.makedirs(app.config['EXTRACTED_FOLDER'])

pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

# Set seed for language detection
DetectorFactory.seed = 0

# Load SpaCy model
nlp = spacy.load('en_core_web_sm')

def extract_text_from_doc(file_path):
    doc = Document(file_path)
    return ' '.join([paragraph.text for paragraph in doc.paragraphs])

def extract_text_from_pdf(file_path):
    text = ''
    pdf_document = fitz.open(file_path)
    for page in pdf_document:
        text += page.get_text()
    return text

def extract_text_from_image(file_path):
    image = Image.open(file_path)
    text = pytesseract.image_to_string(image)
    return text

def translate_text(text, source_lang):
    api_key = 'YourGoogleTranslateAPIKey'  # Replace with your API key
    url = f'https://translation.googleapis.com/language/translate/v2?key={api_key}'
    data = {
        'q': text,
        'source': source_lang,
        'target': 'en',
        'format': 'text'
    }
    response = requests.post(url, data=data)
    translated_text = response.json()['data']['translations'][0]['translatedText']
    return translated_text

def detect_language(text):
    try:
        lang_code = detect(text)
        return lang_code
    except Exception as e:
        print(f"Language detection failed: {e}")
        return None

def save_text_to_docx(text, file_name, folder_path):
    file_path = os.path.join(folder_path, file_name + '_extracted.docx')
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_path)

def preprocess_text(text):
    doc = nlp(text.lower())
    tokens = [token.lemma_ for token in doc if not token.is_stop]
    return ' '.join(tokens)

def compute_similarity_with_transformers(text1, text2):
    model = SentenceTransformer('all-MiniLM-L6-v2')
    embeddings = model.encode([text1, text2])
    return util.cos_sim(embeddings[0], embeddings[1]).item()

def match_requirements_with_resume(job_desc_text, resume_text):
    job_desc_processed = preprocess_text(job_desc_text)
    resume_processed = preprocess_text(resume_text)
    similarity_score = compute_similarity_with_transformers(job_desc_processed, resume_processed)
    return similarity_score

def process_file(file_path, folder_path):
    ext = file_path.split('.')[-1].lower()
    file_name = os.path.splitext(os.path.basename(file_path))[0]

    if ext in ['pdf']:
        text = extract_text_from_pdf(file_path)
    elif ext in ['png', 'jpg', 'jpeg']:
        text = extract_text_from_image(file_path)
    elif ext == 'docx':
        text = extract_text_from_doc(file_path)
    else:
        print(f"Unsupported file type: {ext}")
        return None

    detected_lang = detect_language(text)
    if detected_lang in ['az', 'tr', 'ru']:
        text = translate_text(text, detected_lang)

    save_text_to_docx(text, file_name, folder_path)
    return text

def load_processed_records():
    if not os.path.exists(app.config['PROCESSED_RECORD']):
        return {}
    with open(app.config['PROCESSED_RECORD'], 'r') as file:
        return json.load(file)

def save_processed_records(records):
    with open(app.config['PROCESSED_RECORD'], 'w') as file:
        json.dump(records, file)

@app.route('/upload', methods=['POST'])
@swag_from({
    'summary': 'Upload job requirements and a zip file containing multiple resume files',
    'consumes': ['multipart/form-data'],
    'parameters': [
        {
            'name': 'requirements',
            'in': 'formData',
            'type': 'file',
            'required': True,
            'description': 'Job requirements file (DOCX, PDF, Image)'
        },
        {
            'name': 'resumes_zip',
            'in': 'formData',
            'type': 'file',
            'required': True,
            'description': 'Zip file containing multiple resume files'
        }
    ],
    'responses': {
        200: {
            'description': 'Files processed and similarity scores calculated',
            'schema': {
                'type': 'object',
                'properties': {
                    'results': {
                        'type': 'array',
                        'items': {
                            'type': 'object',
                            'properties': {
                                'file_name': {'type': 'string'},
                                'similarity_score': {'type': 'string'},
                                'original_link': {'type': 'string'}
                            }
                        }
                    }
                }
            }
        },
        400: {
            'description': 'No files part or unsupported file type'
        }
    }
})
def upload_files():
    if 'requirements' not in request.files or 'resumes_zip' not in request.files:
        return jsonify({"msg": "Requirements file and resumes zip file are required"}), 400

    requirements_file = request.files['requirements']
    requirements_filename = secure_filename(requirements_file.filename)
    requirements_file_path = os.path.join(app.config['UPLOAD_FOLDER'], requirements_filename)
    requirements_file.save(requirements_file_path)

    job_desc_text = process_file(requirements_file_path, app.config['UPLOAD_FOLDER'])
    if job_desc_text is None:
        return jsonify({"msg": "Unsupported requirements file type"}), 400

    resumes_zip_file = request.files['resumes_zip']
    resumes_zip_filename = secure_filename(resumes_zip_file.filename)
    resumes_zip_file_path = os.path.join(app.config['ZIP_FOLDER'], resumes_zip_filename)
    resumes_zip_file.save(resumes_zip_file_path)

    job_folder = os.path.join(app.config['UPLOAD_FOLDER'], requirements_filename.split('.')[0])
    if not os.path.exists(job_folder):
        os.makedirs(job_folder)

    processed_records = load_processed_records()
    if requirements_filename not in processed_records:
        processed_records[requirements_filename] = []

    # Extract resumes from the zip file
    with zipfile.ZipFile(resumes_zip_file_path, 'r') as zip_ref:
        zip_ref.extractall(job_folder)

    # Process each extracted file
    results = []
    for root, dirs, files in os.walk(job_folder):
        for file in files:
            if file in processed_records[requirements_filename]:
                continue  # Skip already processed files

            file_path = os.path.join(root, file)
            text = process_file(file_path, app.config['EXTRACTED_FOLDER'])
            if text is None:
                continue

            similarity_score = match_requirements_with_resume(job_desc_text, text)
            # Adjust path to include the subdirectory
            original_link = request.host_url + 'uploads/' + requirements_filename.split('.')[0] + '/' + file
            results.append({
                "file_name": file,
                "similarity_score": f"{similarity_score * 100:.2f}%",
                "original_link": original_link
            })
            processed_records[requirements_filename].append(file)

    save_processed_records(processed_records)

    # Sort results by similarity score in descending order
    results = sorted(results, key=lambda x: float(x['similarity_score'].strip('%')), reverse=True)

    return jsonify({"results": results}), 200

@app.route('/uploads/<path:filename>')
def download_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)